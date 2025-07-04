from docx import Document
from docx.shared import Inches

# Mapping of headings to image filenames
image_map = {
    'System Flowchart': 'flowchart.png',
    'Class Diagram': 'class_diagram.png',
    'Sequence Diagram': 'sequence_diagram.png',
    'Entity-Relationship Diagram': 'er_diagram.png',
}

doc_path = 'Student_Guider_Chatbot_Documentation.docx'
doc = Document(doc_path)

# Collect all paragraphs as text for easier matching
paras = doc.paragraphs

for i, para in enumerate(paras):
    if para.text in image_map:
        # Look ahead for the placeholder and remove it
        for j in range(i+1, min(i+5, len(paras))):
            if '[Insert diagram image here.' in paras[j].text:
                paras[j].clear()
                # Insert image after the description (which is after heading)
                run = paras[j-1].add_run()
                run.add_picture(image_map[para.text], width=Inches(5.5))
                break

# Save the updated document
doc.save('Student_Guider_Chatbot_Documentation_with_Diagrams.docx')
print('Diagrams inserted. Saved as Student_Guider_Chatbot_Documentation_with_Diagrams.docx') 