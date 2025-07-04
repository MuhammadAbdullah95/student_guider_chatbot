import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()
doc.add_heading('Student Guider Chatbot Documentation', 0)

doc.add_heading('1. Project Overview', level=1)
doc.add_paragraph(
    'The Student Guider Chatbot is an AI-powered assistant that helps students with information about studying abroad, including programs, scholarships, eligibility, and application processes. The system consists of a Python backend (API and agent logic) and a modern React frontend (chat UI).'
)

doc.add_heading('2. System Architecture', level=1)
doc.add_paragraph('The system is composed of:')
doc.add_paragraph('- React Frontend (study-buddy-chat-ui)\n- FastAPI Backend (main.py)\n- ChromaDB Vector Store\n- Google Search Tool (fallback)')

# Add diagrams as placeholders
for title, desc in [
    ("System Flowchart", "Shows the flow from user to frontend, backend, vector store, and back."),
    ("Class Diagram", "Shows main backend and frontend classes/components and their relationships."),
    ("Sequence Diagram", "Step-by-step message flow from user to assistant and back."),
    ("Entity-Relationship Diagram", "Session/message structure.")
]:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)
    doc.add_paragraph('[Insert diagram image here. If using Mermaid, export as PNG and replace this placeholder.]')

doc.add_heading('3. Backend (Python/FastAPI)', level=1)
doc.add_paragraph('''
Features:
- Conversational AI agent for study abroad guidance
- Retrieval-augmented generation using a vector store (ChromaDB)
- Live internet search fallback for up-to-date information
- Session-based chat with in-memory session management
- REST API built with FastAPI

API Endpoints:
- POST /chat: Main chat endpoint. Accepts { session_id, message } and returns { session_id, response }
- GET /: Welcome and health check
- GET /health: Service health check
- OPTIONS /chat: CORS preflight

Agent Logic:
- Uses Google Gemini API for embeddings and content generation
- Retrieves relevant context from ChromaDB using semantic search
- Follows strict behavioral guidelines (see instructions.py), including asking clarifying questions, using vector store as primary knowledge, with internet search as fallback, explaining academic terms, and maintaining a supportive, student-friendly tone

Session Management:
- Sessions are tracked in-memory (dictionary keyed by session_id)
- Each session stores a list of message objects (role/content)
''')

doc.add_heading('4. Frontend (React/Vite)', level=1)
doc.add_paragraph('''
Features:
- Modern chat interface for interacting with the assistant
- Session persistence using sessionStorage
- Auto-scrolling, typing indicator, and error handling for smooth UX
- Reusable UI components (buttons, textareas, toasts, tooltips, etc.)
- React Router for page navigation

Main Components & Structure:
- src/App.tsx: Main app entry, sets up providers and routes
- src/pages/Index.tsx: Main chat page, handles message state, API calls, and UI
- src/pages/NotFound.tsx: 404 page
- src/components/: Chat message and typing indicator components
- src/components/ui/: Large library of reusable UI primitives (button, toast, dialog, etc.)
- src/hooks/: Custom React hooks (e.g., for toast notifications)
- src/lib/utils.ts: Utility functions

UI/UX Flow:
1. User opens the app and sees a greeting
2. User types a question and presses Enter
3. The message is sent to the backend; a typing indicator is shown
4. The assistant's response appears in the chat
5. User can start a new session at any time
''')

doc.add_heading('5. Setup & Running', level=1)
doc.add_paragraph('''
Backend:
1. Install dependencies: pip install -r requirements.txt
2. Set up environment variables (e.g., GOOGLE_API_KEY)
3. Run the FastAPI server: uvicorn main:app --reload

Frontend:
1. Navigate to study-buddy-chat-ui/
2. Install dependencies: npm install
3. Start the development server: npm run dev
''')

doc.add_heading('6. Extending & Customizing', level=1)
doc.add_paragraph('''
- To add new knowledge: Update the vector store using set_vector_store.py or similar scripts
- To change agent behavior: Edit instructions.py
- To add new UI features: Create new components in src/components/ or src/components/ui/
''')

doc.add_heading('7. Appendix', level=1)
doc.add_paragraph('''
Python Requirements:
- streamlit
- httpx

Frontend Stack:
- Vite
- TypeScript
- React
- shadcn-ui
- Tailwind CSS

Project Metadata:
- Author: MuhammadAbdullah95 (ma2404374@gmail.com)
- Python >= 3.11
- Version: 0.1.0
''')

# Save the document
doc.save('Student_Guider_Chatbot_Documentation.docx')
print('Documentation generated: Student_Guider_Chatbot_Documentation.docx')
